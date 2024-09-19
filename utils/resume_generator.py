import json
import os

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from . import page_dimensions as page


def generate_resume(profile_data, output_file):
    create_resume_from_json(profile_data, output_file)


def create_resume_from_json(data, output_file):
    doc = Document()
    for key, value in data["Personal Info"].items():
        if key == "Name":
            name_paragraph = doc.add_paragraph()
            name_run = name_paragraph.add_run(value)
            name_run.bold = True
            name_run.font.size = Inches(0.4)
            student = name_paragraph.add_run("    Student")
            student.italic = True
            student.font.size = Inches(0.2)
            continue
        if key == "Date of Birth":
            continue
        if key == "Address":
            continue
    if "social_media" in data:
        add_social_media_icons(
            doc, data["social_media"][0], os.path.join("app", "icons")
        )
    if "Profile Summary" in data and data["Profile Summary"] is not None:
        paragraph = doc.add_paragraph()
        header_paragraph = paragraph.add_run("Profile")
        header_paragraph.font.size = Pt(16)
        header_paragraph.bold = True
        header_paragraph.underline = False
        header_paragraph.style.font.color.rgb = RGBColor(0, 0, 0)
        name_run = paragraph.add_run(
            " ____________________________________________________________________________________________________"
        )
        name_run.bold = True
        name_run.font.size = Inches(0.2)
        doc.add_paragraph(f"{data['About_me']}")
    paragraph = doc.add_paragraph()
    header_paragraph = paragraph.add_run("Skills")
    header_paragraph.font.size = Pt(16)
    header_paragraph.bold = True
    header_paragraph.underline = False
    header_paragraph.style.font.color.rgb = RGBColor(0, 0, 0)
    name_run = paragraph.add_run(
        " ___________________________________________________________________________________________________"
    )
    name_run.bold = True
    name_run.font.size = Inches(0.2)
    for skill in data["Skills"]:
        doc.add_paragraph(skill)
    paragraph = doc.add_paragraph()
    header_paragraph = paragraph.add_run("Education")
    header_paragraph.font.size = Pt(16)
    header_paragraph.bold = True
    header_paragraph.underline = False
    header_paragraph.style.font.color.rgb = RGBColor(0, 0, 0)
    name_run = paragraph.add_run(
        " ________________________________________________________________________________________________"
    )
    name_run.bold = True
    name_run.font.size = Inches(0.2)
    for education in data["education"]:
        if (
            "timePeriod" in education
            and "startDate" in education["timePeriod"]
            and "year" in education["timePeriod"]["startDate"]
        ):
            start_year = education["timePeriod"]["startDate"]["year"]
            start = f" {start_year} - "
        else:
            start = ""
        if (
            "timePeriod" in education
            and "endDate" in education["timePeriod"]
            and "year" in education["timePeriod"]["endDate"]
        ):
            end_year = education["timePeriod"]["endDate"]["year"]
        else:
            end_year = "Present "
        if "degreeName" in education:
            degree_name = education.get("degreeName", "")
            degree = f"{degree_name} in "
        else:
            degree = ""
        if "fieldOfStudy" in education:
            field_of_study = education.get("fieldOfStudy", "")
            field = f"{field_of_study}, \n                          "
        else:
            field = ""
        if "schoolName" in education:
            school_name = education.get("schoolName", "")
            school = f"{school_name}"
        else:
            school = ""
        doc.add_paragraph(f"{start}{end_year}   {degree}{field}{school}")
    paragraph = doc.add_paragraph()
    header_paragraph = paragraph.add_run("Project")
    header_paragraph.font.size = Pt(16)
    header_paragraph.bold = True
    header_paragraph.underline = False
    header_paragraph.style.font.color.rgb = RGBColor(0, 0, 0)
    name_run = paragraph.add_run(
        " ____________________________________________________________________________________________________"
    )
    name_run.bold = True
    name_run.font.size = Inches(0.2)
    paragraph = doc.add_paragraph()
    for projects in data["Projects"]:
        if projects["Language"] != None:
            name_run = paragraph.add_run(
                f"{projects['Name']}  ({projects['Language']}) \n  "
            )
            name_run.bold = True
            if "Description" in projects:
                des_run = paragraph.add_run(f"{projects['Description']} \n \n")
                des_run.italic = True
        else:
            name_run = paragraph.add_run(f"{projects['Name']} \n  ")
            name_run.bold = True
            if "Description" in projects:
                des_run = paragraph.add_run(f"{projects['Description']} \n \n")
                des_run.italic = True
    paragraph = doc.add_paragraph()
    header_paragraph = paragraph.add_run("Certificates")
    header_paragraph.font.size = Pt(16)
    header_paragraph.bold = True
    header_paragraph.underline = False
    header_paragraph.style.font.color.rgb = RGBColor(0, 0, 0)
    name_run = paragraph.add_run(
        " _________________________________________________________________________________________________"
    )
    name_run.bold = True
    name_run.font.size = Inches(0.2)
    paragraph = doc.add_paragraph()
    for certificate in data["certificates"]:
        name_run = paragraph.add_run(f"{certificate['name']}\n")
        name_run.bold = True
        if "authority" in certificate:
            des_run = paragraph.add_run(f"{certificate['authority']}\n")
            des_run.italic = True
    page.set_page_height(doc, 15)
    page.set_page_width(doc, 10)
    page.set_page_left_margin(doc, 0.5)
    page.set_page_right_margin(doc, 0.5)
    page.set_page_top_margin(doc, 0.5)
    page.set_page_bottom_margin(doc, 0.5)
    doc.save(output_file)


def add_horizontal_line(doc):
    """
    Add a horizontal line to the document.
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run()
    run.add_break()
    run.add_picture(
        os.path.join("app", "icons", "profile1.png"),
        width=Inches(6.9),
        height=Inches(0.5),
    )
    paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE


def count_entries(json_file, section_name):
    with open(json_file, "r") as f:
        data = json.load(f)
        entries = data.get(section_name)
        if entries:
            return len(entries)
        else:
            return 0


def add_icon_with_username(cell, icon_path, username):
    """
    Add an icon with its username below it in a table cell.
    """
    run = cell.paragraphs[0].add_run()
    run.add_picture(icon_path, width=Inches(0.25), height=Inches(0.25))
    cell.add_paragraph(username, style="Quote")


def add_social_media_icons(doc, social_media_data, icons_folder):
    """
    Add social media icons with hyperlinks to the document.
    """
    count = 0
    table = doc.add_table(rows=2, cols=4)
    for social_media in social_media_data:
        social_media_name = social_media
        if social_media_name:
            icon_filename = social_media_name.lower().replace(" ", "_") + ".png"
            icon_path = os.path.join(icons_folder, icon_filename)
            if os.path.exists(icon_path):
                cell = table.cell(0, count)
                add_icon_with_username(
                    cell, icon_path, social_media_data[social_media_name]
                )
                count += 1
            else:
                print(f"Icon file not found for {social_media_name}: {icon_path}")
        else:
            print("Social media name not specified")
